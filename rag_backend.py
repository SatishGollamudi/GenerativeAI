"""
Document Q&A RAG System - Backend
Uses Groq for LLM inference and sentence-transformers for embeddings.

Install dependencies:
    pip install groq sentence-transformers faiss-cpu PyPDF2 numpy python-dotenv python-docx

Usage:
    python rag_backend.py
"""

import os
import sys
import hashlib
import numpy as np
from pathlib import Path
from dataclasses import dataclass

# Fix Windows CMD encoding - write directly to buffer, no TextIOWrapper
if sys.platform == "win32":
    import ctypes
    ctypes.windll.kernel32.SetConsoleOutputCP(65001)
    os.environ["PYTHONIOENCODING"] = "utf-8"

def safe_print(text: str) -> None:
    """Write text directly to stdout buffer as UTF-8 — bypasses all encoding layers."""
    sys.stdout.buffer.write((text + "\n").encode("utf-8", errors="replace"))
    sys.stdout.buffer.flush()

try:
    import faiss
except ImportError:
    raise ImportError("Run: pip install faiss-cpu")

try:
    from groq import Groq
except ImportError:
    raise ImportError("Run: pip install groq")

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    raise ImportError("Run: pip install sentence-transformers")

try:
    import PyPDF2
except ImportError:
    raise ImportError("Run: pip install PyPDF2")

try:
    import docx as python_docx
except ImportError:
    raise ImportError("Run: pip install python-docx")

from dotenv import load_dotenv
load_dotenv()


# ---------------------------------------------
#  Data Structures
# ---------------------------------------------

@dataclass
class Chunk:
    text: str
    source: str
    page: int
    chunk_id: str = ""

    def __post_init__(self):
        self.chunk_id = hashlib.md5(
            f"{self.source}:{self.page}:{self.text[:50]}".encode("utf-8")
        ).hexdigest()[:8]


@dataclass
class SearchResult:
    chunk: Chunk
    score: float


@dataclass
class RAGResponse:
    answer: str
    sources: list[SearchResult]
    query: str


# ---------------------------------------------
#  Document Loader
# ---------------------------------------------

class DocumentLoader:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load(self, path: str) -> list[Chunk]:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {path}")
        ext = p.suffix.lower()
        if ext == ".pdf":
            return self._load_pdf(p)
        elif ext == ".docx":
            return self._load_docx(p)
        else:
            return self._load_text(p)

    def _load_pdf(self, path: Path) -> list[Chunk]:
        chunks = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                for chunk in self._split(text):
                    chunks.append(Chunk(text=chunk, source=path.name, page=page_num + 1))
        return chunks

    def _load_docx(self, path: Path) -> list[Chunk]:
        doc = python_docx.Document(str(path))
        full_text = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text += "\n" + cell.text.strip()
        return [
            Chunk(text=chunk, source=path.name, page=i + 1)
            for i, chunk in enumerate(self._split(full_text))
        ]

    def _load_text(self, path: Path) -> list[Chunk]:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [
            Chunk(text=chunk, source=path.name, page=i + 1)
            for i, chunk in enumerate(self._split(text))
        ]

    def _split(self, text: str) -> list[str]:
        words = text.split()
        results, i = [], 0
        while i < len(words):
            chunk = " ".join(words[i: i + self.chunk_size])
            if chunk.strip():
                results.append(chunk)
            i += self.chunk_size - self.chunk_overlap
        return results


# ---------------------------------------------
#  Vector Store (FAISS)
# ---------------------------------------------

class VectorStore:
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        safe_print(f"[VectorStore] Loading embedding model '{model_name}' ...")
        self.model = SentenceTransformer(model_name)
        self.dim = self.model.get_sentence_embedding_dimension()
        self.index = faiss.IndexFlatIP(self.dim)
        self.chunks: list[Chunk] = []

    def add(self, chunks: list[Chunk]) -> None:
        if not chunks:
            return
        embeddings = self._encode([c.text for c in chunks])
        self.index.add(embeddings)
        self.chunks.extend(chunks)
        safe_print(f"[VectorStore] Indexed {len(chunks)} chunks  (total {len(self.chunks)})")

    def search(self, query: str, k: int = 5) -> list[SearchResult]:
        if self.index.ntotal == 0:
            return []
        q_emb = self._encode([query])
        scores, indices = self.index.search(q_emb, min(k, self.index.ntotal))
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx >= 0:
                results.append(SearchResult(chunk=self.chunks[idx], score=float(score)))
        return results

    def clear(self) -> None:
        self.index.reset()
        self.chunks.clear()

    def _encode(self, texts: list[str]) -> np.ndarray:
        vecs = self.model.encode(texts, normalize_embeddings=True, show_progress_bar=False)
        return vecs.astype(np.float32)


# ---------------------------------------------
#  RAG Engine
# ---------------------------------------------

class RAGEngine:
    SYSTEM_PROMPT = (
        "You are a precise document assistant. "
        "Answer the user's question using ONLY the provided context excerpts. "
        "If the context does not contain enough information, say so honestly. "
        "Always cite which document(s) you drew information from."
    )

    def __init__(
        self,
        groq_model: str = "llama-3.3-70b-versatile",
        embed_model: str = "all-MiniLM-L6-v2",
        top_k: int = 5,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY not set. Add it to your .env file.")

        self.client     = Groq(api_key=api_key)
        self.groq_model = groq_model
        self.top_k      = top_k
        self.loader     = DocumentLoader(chunk_size, chunk_overlap)
        self.store      = VectorStore(embed_model)
        self.loaded_docs: list[str] = []

    def ingest(self, path: str) -> int:
        chunks = self.loader.load(path)
        self.store.add(chunks)
        self.loaded_docs.append(Path(path).name)
        return len(chunks)

    def query(self, question: str) -> RAGResponse:
        results = self.store.search(question, k=self.top_k)
        context = self._build_context(results)
        answer  = self._generate(question, context)
        return RAGResponse(answer=answer, sources=results, query=question)

    def clear(self) -> None:
        self.store.clear()
        self.loaded_docs.clear()

    def _build_context(self, results: list[SearchResult]) -> str:
        parts = []
        for r in results:
            parts.append(
                f"[Source: {r.chunk.source}, page {r.chunk.page}]\n{r.chunk.text}"
            )
        return "\n\n---\n\n".join(parts)

    def _generate(self, question: str, context: str) -> str:
        user_msg = (
            f"Context excerpts:\n\n{context}\n\n"
            f"Question: {question}\n\n"
            "Answer:"
        )
        completion = self.client.chat.completions.create(
            model=self.groq_model,
            messages=[
                {"role": "system", "content": self.SYSTEM_PROMPT},
                {"role": "user",   "content": user_msg},
            ],
            temperature=0.2,
            max_tokens=1024,
        )
        return completion.choices[0].message.content


# ---------------------------------------------
#  Save to answer.txt with UTF-8 BOM for Notepad
# ---------------------------------------------

def save_and_open(response: RAGResponse) -> None:
    out_path = Path("answer.txt")

    lines = [
        "=" * 60,
        f"  QUESTION:  {response.query}",
        "=" * 60,
        "",
        "ANSWER:",
        "",
        response.answer,
        "",
        "-" * 60,
        "SOURCES USED:",
    ]
    for r in response.sources:
        lines.append(
            f"  - {r.chunk.source}  (page {r.chunk.page})  similarity={r.score:.3f}"
        )
    lines.append("-" * 60)

    # Write raw UTF-8 bytes with BOM directly — no encoding layer issues
    content = "\r\n".join(lines)
    with open(out_path, "wb") as f:
        f.write(b"\xef\xbb\xbf")  # UTF-8 BOM
        f.write(content.encode("utf-8", errors="replace"))

    safe_print("\n  Answer saved to 'answer.txt'  ->  opening in Notepad ...\n")
    os.system("notepad answer.txt")


# ---------------------------------------------
#  CLI
# ---------------------------------------------

def main():
    safe_print("=" * 60)
    safe_print("  Document Q&A RAG System  (Groq + FAISS)")
    safe_print("=" * 60)

    rag = RAGEngine()

    while True:
        sys.stdout.buffer.write(b"\nEnter document path (PDF, TXT, DOCX) or press Enter to skip: ")
        sys.stdout.buffer.flush()
        path = input().strip()
        if not path:
            break
        try:
            n = rag.ingest(path)
            safe_print(f"  OK  Indexed {n} chunks from '{Path(path).name}'")
        except Exception as e:
            safe_print(f"  ERROR  {e}")

    if not rag.loaded_docs:
        safe_print("\nNo documents loaded. Exiting.")
        return

    safe_print(f"\nDocuments loaded: {', '.join(rag.loaded_docs)}")
    safe_print("Type 'quit' to exit.\n")

    while True:
        sys.stdout.buffer.write(b"You: ")
        sys.stdout.buffer.flush()
        question = input().strip()

        if not question or question.lower() in {"quit", "exit", "q"}:
            safe_print("\nGoodbye!")
            break

        try:
            response = rag.query(question)
            save_and_open(response)

            safe_print("Sources used:")
            for r in response.sources:
                safe_print(f"  * {r.chunk.source} (page {r.chunk.page})  similarity={r.score:.3f}")
            safe_print("")

        except Exception as e:
            safe_print(f"\n  ERROR  {e}\n")


if __name__ == "__main__":
    main()